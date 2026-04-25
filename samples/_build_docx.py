"""One-off helper to render the sample .md files as .docx.

Run from the project root:
    python samples/_build_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent

SOURCES = {
    "company_handbook.md": "company_handbook.docx",
    "installation_guide.md": "installation_guide.docx",
}


def render(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    table_rows: list[list[str]] | None = None

    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()

        if line.startswith("|") and "|" in line[1:]:
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(set(c) <= set("-: ") for c in cells):
                continue  # markdown table separator
            if table_rows is None:
                table_rows = []
            table_rows.append(cells)
            continue

        if table_rows is not None:
            _emit_table(doc, table_rows)
            table_rows = None

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line), style="List Number")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(_strip_inline(line))

    if table_rows is not None:
        _emit_table(doc, table_rows)

    doc.save(docx_path)


def _strip_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def _emit_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Light Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(width):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = _strip_inline(row[c_idx]) if c_idx < len(row) else ""


def main() -> None:
    for src, dst in SOURCES.items():
        render(ROOT / src, ROOT / dst)
        print(f"wrote {dst}")


if __name__ == "__main__":
    main()
